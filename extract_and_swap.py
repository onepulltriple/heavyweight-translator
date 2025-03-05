# Function definitions for extraction, deduplication, and swapping
import docx
import constants 
import file_paths as FP 
from docx import Document
from auxiliary_operations import *
from conditions_checks import *
from dict_operations import *
from csv_read_operations import *
from csv_write_operations import *
from xml_operations import *
from preservation_operations import *
from progress_indication_operations import *
from tagging_operations import *
from copy import deepcopy
import math
from xml.sax.saxutils import escape, unescape


#__________________________________________________________________________
###########################################################################
# Function to extract or swap text elements from a docx file
# Argument ordering for functions: translation_dict, paragraph, current_run, variables/counters
def extract_or_swap_text_in_docx(input_file, step, translation_dict = {}, output_docx = None):

    # Read the unmodified input .docx document into memory
    doc = Document(input_file)
    
    # Initialize operation counters
    percentage_increment_to_report = 1 #percent

    if step == constants.EXTRACT:
        count_of_relevant_paragraphs = count_paragraphs(doc)
        # A -1 should be returned when no extraction occurs for a paragraph, 
        # thereby reducing the expected total number of paragraphs to process
        newest_print_progress_threshold = math.ceil(percentage_increment_to_report/100*count_of_relevant_paragraphs)
    
    if step == constants.SWAP:
        # A +1 should be returned for all successful swaps
        # thereby indicating progress towards swapping of relevant paragraphs
        count_of_relevant_paragraphs = 0
        newest_print_progress_threshold = math.ceil(percentage_increment_to_report/100*len(translation_dict))
    
    print_progress_increment = newest_print_progress_threshold

    # PARAGRAPHS ##########################################################
    for paragraph in doc.paragraphs:
        count_of_relevant_paragraphs += process_paragraph_and_runs_within_it(translation_dict, paragraph, step) #add doc if debugging is needed
        newest_print_progress_threshold = indicate_progress(translation_dict, step, newest_print_progress_threshold, print_progress_increment, count_of_relevant_paragraphs)
        

    # TABLES ##############################################################
    for table in doc.tables:
        count_of_relevant_paragraphs = process_table_cells(translation_dict, table, step, newest_print_progress_threshold, print_progress_increment, count_of_relevant_paragraphs) #add doc if debugging is needed
        newest_print_progress_threshold = indicate_progress(translation_dict, step, newest_print_progress_threshold, print_progress_increment, count_of_relevant_paragraphs)


    # RESULTS #############################################################
    if step == constants.EXTRACT:
        write_dict_to_json(translation_dict, FP.TEMP_translation_dict_file_path)
        #write_translation_dict_to_csv(translation_dict, FP.source_language_plain_texts_file_path)
        write_translation_dict_to_csv_simplified(translation_dict, FP.source_language_plain_texts_file_path)
        print(f"There were {len(translation_dict)} {step} operations.\n")
        
    if step == constants.SWAP:
        print(f"There were {count_of_relevant_paragraphs} {step} operations.\n")
        #print(f"{total_no_swap_count} {step} operations failed.\n")
        # Save the modified document to the output file
        print("Saving translated document...\n")
        doc.save(output_docx)
    
    print("Done.\n")

#__________________________________________________________________________
###########################################################################
# Function to consolidate the text in runs from a paragraph and its hyperlinks
def consolidate_runs(paragraph): #add doc if debugging is needed

    # Prepare to consolidate text-having runs
    text_consolidator = ""

    # Hold on to last run when needed
    previous_run = None

    index_of_run = -1

    # Loop over all the runs/hyperlinks in the paragraph
    for current_run_or_hyperlink, next_run_or_hyperlink in pairwise_circular(paragraph.iter_inner_content()):
        index_of_run += 1

        # Skip pictures or other non-text-having objects
        if not current_run_or_hyperlink.text:
            #current_run_or_hyperlink.text = glyph_tag(index_of_run)
            previous_run = current_run_or_hyperlink
            continue

        # If the current object is a hyperlink
        if isinstance(current_run_or_hyperlink, docx.text.hyperlink.Hyperlink):
            # Rename object for clarity
            current_hyperlink = current_run_or_hyperlink
            # Don't consolidate, just deal with the text from each text-having run

            # This treatment could perhaps be simplified. Maybe hyperlinks don't need consolidation (to be verified)
            for current_run in current_hyperlink.runs:
                if current_run.text:
                    # collect text
                    text_consolidator += current_run.text
                    # dump consolidator into this run
                    current_run.text = text_consolidator                   
                    # clear the text consolidator and move on
                    text_consolidator = ""

            previous_run = current_run_or_hyperlink
            continue

        # Otherwise, if the current object is a run    
        elif isinstance(current_run_or_hyperlink, docx.text.run.Run):
            # Rename object for clarity and to avoid mutating the current_run_or_hyperlink
            current_run = current_run_or_hyperlink

            # Conditions under which to collect and dump immediately
            if (the_current_run_has_an_R_character(current_run)
                or button_like_formatting_starts_and_ends_in_the_current_run(current_run, text_consolidator)
                or button_like_formatting_starts_and_ends_in_the_next_run(next_run_or_hyperlink, text_consolidator)
                or weird_symbol_bracketed_by_blank_char_starts_in_the_current_run(current_run, next_run_or_hyperlink, text_consolidator) 
                or weird_symbol_bracketed_by_blank_char_ends_in_the_current_run(previous_run, current_run, next_run_or_hyperlink, text_consolidator) 
                ):
                # collect text
                text_consolidator += current_run.text
                # dump consolidator into this run
                current_run.text = text_consolidator                   
                # clear the text consolidator and move on
                text_consolidator = ""
                previous_run = current_run_or_hyperlink
                continue
            
            # Conditions under which to keep consolidating
            if (either_has_special_characters(current_run, next_run_or_hyperlink)
                and (the_current_run_has_one_or_two_special_characters(current_run)
                     or the_next_run_has_one_or_two_special_characters(next_run_or_hyperlink))
                ):
                # collect this run's text
                text_consolidator += current_run.text
                # clear this run's text
                current_run.text = ignore_run_tag(index_of_run)
                
                if the_last_run_in_the_paragraph_has_been_reached(next_run_or_hyperlink):
                    pass # move on to the code below
                else:
                    previous_run = current_run_or_hyperlink
                    continue # keep consolidating

            if (button_like_formatting_starts_in_this_run(current_run, next_run_or_hyperlink, text_consolidator)  
                or button_like_formatting_ends_in_the_next_run(next_run_or_hyperlink, text_consolidator)         
                ):
                # collect this run's text
                text_consolidator += current_run.text
                # clear this run's text
                current_run.text = ignore_run_tag(index_of_run)
                
                if the_last_run_in_the_paragraph_has_been_reached(next_run_or_hyperlink):
                    pass # move on to the code below
                else:
                    previous_run = current_run_or_hyperlink
                    continue # keep consolidating

            if (bogus_change_of_nature_conditions_are_found(previous_run, current_run, next_run_or_hyperlink)         
                ):
                # collect this run's text
                text_consolidator += current_run.text
                # clear this run's text
                current_run.text = ignore_run_tag(index_of_run)
                
                if the_last_run_in_the_paragraph_has_been_reached(next_run_or_hyperlink):
                    pass # move on to the code below
                else:
                    previous_run = current_run_or_hyperlink
                    continue # keep consolidating

            # If one of these conditions is met, stop consolidating
            if (the_last_run_in_the_paragraph_has_been_reached(next_run_or_hyperlink)
                  or there_is_no_text_in_the_next_run(next_run_or_hyperlink)
                  or internal_hidden_text_style_has_been_reached(next_run_or_hyperlink)
                  or button_like_formatting_starts_in_next_run(next_run_or_hyperlink, text_consolidator)
                  or button_like_formatting_ends_in_this_run(current_run, text_consolidator)
                  or there_is_a_change_of_nature(current_run, next_run_or_hyperlink)
                ): 
                # collect text
                text_consolidator += current_run.text
                # dump consolidator into this run
                current_run.text = text_consolidator
                # clear the text collector and move on
                text_consolidator = ""                
                previous_run = current_run_or_hyperlink
                continue

            print(f"The text-having run \"{current_run_or_hyperlink.text}\" was not handled.\n")

    return paragraph

#__________________________________________________________________________
###########################################################################
# Function to extract and tag text from a paragraph and its hyperlinks
# Returns a string with tagged runs intermixed with plain text
def extract_runs(paragraph_with_cons_runs):

    index_of_run = -1
    paragraph_tagged_source_text_with_preserves = ""
    cons_run_tagged_text_with_preserves = ""

    # Hold on to last run when needed
    previous_run = None

    # Loop over all the runs/hyperlinks in the paragraph
    for current_run_or_hyperlink in paragraph_with_cons_runs.iter_inner_content():
        index_of_run += 1

        # If the object is a non-text-having run, e.g. pictures or symbols (which might also be pictures)
        if not current_run_or_hyperlink.text:
            # Rename object for clarity
            current_glyph_holder = current_run_or_hyperlink
            # Assign a placeholder to the object's text field
            current_glyph_holder.text = glyph_tag(index_of_run)
            # There is nothing to preserve
            cons_run_plain_text_with_preserves = current_glyph_holder.text
            # Do not add a tag
            cons_run_tagged_text_with_preserves = cons_run_plain_text_with_preserves
            # (style would be Default Paragraph Font)
            cons_run_style = current_glyph_holder.style.name
            
        # Otherwise, if the current object is a hyperlink
        elif isinstance(current_run_or_hyperlink, docx.text.hyperlink.Hyperlink):
            # Rename object for clarity
            current_hyperlink = current_run_or_hyperlink
            # Preserve the consolidated run's special characters
            cons_run_plain_text_with_preserves = preserve_run_special_items_with_temp_symbols(escape(current_hyperlink.text))
            # It needs a tag
            cons_run_tagged_text_with_preserves = hyperlink_tag(cons_run_plain_text_with_preserves, index_of_run)
            # (style would be Hyperlink)
            cons_run_style = "Hyperlink"

        # Otherwise, if the current object is a run    
        elif isinstance(current_run_or_hyperlink, docx.text.run.Run):
            # Rename object for clarity
            current_run = current_run_or_hyperlink
            # If the run is of a non-default style
            if (current_run.style.name != "Default Paragraph Font" 
                    # and it is not a cleared run
                    and current_run.text != ignore_run_tag(index_of_run)):
                # Preserve the consolidated run's special characters
                cons_run_plain_text_with_preserves = preserve_run_special_items_with_temp_symbols(escape(current_run.text))
                # Give it a tag
                cons_run_tagged_text_with_preserves = styled_run_tag(cons_run_plain_text_with_preserves, index_of_run)
                # (style would be current_run.style.name)
            elif(the_current_run_has_an_R_character(current_run)
                or there_WAS_a_change_of_nature(current_run, previous_run)
                    # and it is not a cleared run
                    and current_run.text != ignore_run_tag(index_of_run)
                    # and it is not an empty run
                    and not current_run.text.isspace()):
                # Preserve the consolidated run's special characters
                cons_run_plain_text_with_preserves = preserve_run_special_items_with_temp_symbols(escape(current_run.text))
                # Give it a tag
                cons_run_tagged_text_with_preserves = changed_run_tag(cons_run_plain_text_with_preserves, index_of_run)
                # (style would be Default Paragraph Font, probably)
            else: 
                # Preserve the consolidated run's special characters
                cons_run_plain_text_with_preserves = preserve_run_special_items_with_temp_symbols(escape(current_run.text))
                # It does not need a tag (don't add an ignore tag)
                cons_run_tagged_text_with_preserves = cons_run_plain_text_with_preserves
                # (style would be Default Paragraph Font)
            # Get style
            cons_run_style = current_run.style.name
        

        # If the consolidated run is of a non-default style
        if cons_run_style != "Default Paragraph Font":
            # Append it to the paragraph's tagged text with tags
            paragraph_tagged_source_text_with_preserves += cons_run_tagged_text_with_preserves
        else: # The run is of the default style
            if cons_run_tagged_text_with_preserves == changed_run_tag(cons_run_plain_text_with_preserves, index_of_run):
                # Append with tags
                paragraph_tagged_source_text_with_preserves += cons_run_tagged_text_with_preserves
            elif cons_run_plain_text_with_preserves != ignore_run_tag(index_of_run):
                # Append without tags
                paragraph_tagged_source_text_with_preserves += cons_run_plain_text_with_preserves

        previous_run = current_run_or_hyperlink

    return paragraph_tagged_source_text_with_preserves


#__________________________________________________________________________
###########################################################################
# Function to orchestrate the swapping of run-level text for each paragraph
# Returns the translated paragraph and the count of no-swaps (either 1 or 0)
def paragraph_level_swapper(translation_dict, paragraph_with_cons_runs): #add doc if debugging is needed
   
    # An untouched copy of the consolidated paragraph is needed to obtain unchanged info from the consolidated runs
    # Therefore, obtain a carbon copy to give ot the extraction function, which otherwise mutates consolidate paragraphs
    carbon_copy_of_paragraph_with_cons_runs = deepcopy(paragraph_with_cons_runs)

    # Get the tagged version of this paragraph by performing the extraction step again
    paragraph_tagged_source_text_with_preserves = extract_runs(carbon_copy_of_paragraph_with_cons_runs)
    
    # Attempt to find a translation in the dictionary
    if paragraph_tagged_source_text_with_preserves not in translation_dict:
        print(f"The text element \"{paragraph_tagged_source_text_with_preserves}\" was not found in the translation dictionary's keys.")
        
        # Indicate failure
        return paragraph_with_cons_runs, 0
    
    if (paragraph_tagged_source_text_with_preserves != "" 
        and not paragraph_tagged_source_text_with_preserves.isspace()
        #and paragraph_tagged_source_text_with_preserves in translation_dict # comment out if check is already performed above
        ):
        # Get the paragraph's translated counterpart
        paragraph_tagged_translated_text_with_preserves = translation_dict[paragraph_tagged_source_text_with_preserves]['paragraph_tagged_translated_text_with_preserves']

    # Unpreserve the translation pulled from the dictionary
    paragraph_tagged_translated_text = unpreserve_paragraph_translation(paragraph_tagged_translated_text_with_preserves)
    # Break it into objects (dictionaries)
    translated_runs_with_tags = split_string_into_list_of_tagged_and_untagged_elements(paragraph_tagged_translated_text)

    # Swap runs
    translated_paragraph = swap_runs(paragraph_with_cons_runs, translated_runs_with_tags) #add doc if debugging is needed
    
    # Indicate successful swap
    return translated_paragraph, 1


#__________________________________________________________________________
###########################################################################
# Function to swap text in a consolidated paragraph and its hyperlinks on a per-run basis
def swap_runs(paragraph_with_cons_runs, translated_runs_with_tags): #add doc if debugging is needed

    index_of_consolidated_run = -1
    index_of_translated_run = 0

    # An untouched copy of the consolidated paragraph is needed to obtain unchanged info from the consolidated runs
    carbon_copy_of_paragraph_with_cons_runs = list(paragraph_with_cons_runs.iter_inner_content())
    carbon_copy_of_paragraph_with_cons_runs = deepcopy(carbon_copy_of_paragraph_with_cons_runs)

    # Loop over all the runs/hyperlinks in the paragraph with consolidated runs
    for current_run_or_hyperlink in paragraph_with_cons_runs.iter_inner_content():
        index_of_consolidated_run += 1

        # As long as we are still working through the translated runs
        if index_of_translated_run < len(translated_runs_with_tags):      
            # Get the current entry from the list of translated runs
            current_translated_run_dict = translated_runs_with_tags[index_of_translated_run]

            # If the current translated run is a glyph placeholder
            if ("type" in current_translated_run_dict.keys()
                and current_translated_run_dict["type"] == "glyph"):
                # and if the current consolidated run is also a glyph
                if (#"glyph" in current_run_or_hyperlink.text
                    not current_run_or_hyperlink.text 
                    # and it is the correct/target glyph
                    and index_of_consolidated_run == current_translated_run_dict["run_index"]
                    ):
                    # Rename object for clarity
                    current_glyph_holder = current_run_or_hyperlink
                    # Do nothing further (later replace with translated image)
                    
                    index_of_translated_run += 1 # in everycase except when a hyperlink is reached
                # Otherwise, clear this consolidated run and set it to default
                else:
                    clear_cons_run_and_set_to_defaults(current_run_or_hyperlink)

            # Otherwise, if the current translated run is a hyperlink
            elif ("type" in current_translated_run_dict.keys()
                and current_translated_run_dict["type"] == "hyperlink"):
                # and if the current consolidated run is also a hyperlink
                if (isinstance(current_run_or_hyperlink, docx.text.hyperlink.Hyperlink)
                    # and it is the correct/target hyperlink
                    and index_of_consolidated_run == current_translated_run_dict["run_index"]
                    ):
                    # Rename object for clarity
                    current_hyperlink = current_run_or_hyperlink                        
                    # Take the text from the translated hyperlink's text
                    current_hyperlink.runs[0].text = unescape(current_translated_run_dict["text"])
                        # If remaining runs in the hyperlink need to be cleared, do it here
                    # Increment the translated run index
                    index_of_translated_run += 1 
                # Otherwise, clear this consolidated run and set it to default
                else:
                    clear_cons_run_and_set_to_defaults(current_run_or_hyperlink)

            # Otherwise, deal with stylized or plain text runs
            else:
                #if isinstance(current_run_or_hyperlink, docx.text.run.Run):
                # Rename object for clarity
                current_run = current_run_or_hyperlink        
                # Take the text from the translated run
                if "text" in current_translated_run_dict.keys():
                    current_run.text = unescape(current_translated_run_dict["text"])
                else:
                    current_run.text = ""

                # If there is a style applied or changes to retrieve, get and apply
                if "type" in current_translated_run_dict.keys():
                    # Get the index of the consolidated run whose style is needed
                    index_of_styled_run = current_translated_run_dict["run_index"]

                    if (current_translated_run_dict["type"] == "styled"):
                        # Get the consolidated run whose style is needed
                        run_that_has_style_to_apply = carbon_copy_of_paragraph_with_cons_runs[index_of_styled_run]
                        # Apply the style to the current consolidated run
                        current_run.style = run_that_has_style_to_apply.style
                    elif (current_translated_run_dict["type"] == "changed"):
                        # Get the consolidated run whose changes are needed
                        run_that_has_changes_to_apply = carbon_copy_of_paragraph_with_cons_runs[index_of_styled_run]
                        # Apply the changes
                        current_run.font.color.rgb = run_that_has_changes_to_apply.font.color.rgb
                        current_run.font.size = run_that_has_changes_to_apply.font.size
                        current_run.font.name = run_that_has_changes_to_apply.font.name
                        # Apply the style to the current consolidated run
                        current_run.style = run_that_has_changes_to_apply.style

                # Otherwise apply the default style and remove any manually applied changes
                else:
                    # DO NOT set style.name here, as that would rename the style for the whole document
                    current_run.style = "Default Paragraph Font" 
                    current_run.font.color.rgb = current_run._parent.style.font.color.rgb
                    current_run.font.size = None
                    current_run.font.name = None
                    current_run.font.subscript = None

                # Increment the translated run index
                index_of_translated_run += 1 
        
        # Otherwise, there are no more translated runs, so clean up remaining consolidated runs
        else: 
            clear_cons_run_and_set_to_defaults(current_run_or_hyperlink)

    return paragraph_with_cons_runs

#__________________________________________________________________________
###########################################################################
def clear_cons_run_and_set_to_defaults(current_run_or_hyperlink):
    if not current_run_or_hyperlink.text:
        # Rename object for clarity
        current_glyph_holder = current_run_or_hyperlink
        # Do nothing (later replace with translated image)
        print("Glyphs should not make it to this function. How could this have happened?")
        return current_glyph_holder
        
    # Otherwise, if the current object is a hyperlink
    elif isinstance(current_run_or_hyperlink, docx.text.hyperlink.Hyperlink):
        # Rename object for clarity
        current_hyperlink = current_run_or_hyperlink
        print("Hyperlinks should not make it to this function. How could this have happened?")
        return current_hyperlink

    # Otherwise, if the current object is a run    
    elif isinstance(current_run_or_hyperlink, docx.text.run.Run):
        # Rename object for clarity
        current_run = current_run_or_hyperlink        
        # Take the text from the translated run
        current_run.clear()
        # Apply the default style
        # DO NOT set style.name here, as that would rename the style for the whole document
        current_run.style = "Default Paragraph Font" 
        return current_run


#__________________________________________________________________________
###########################################################################
# Function to process table cells and any further nested tables, etc.
def process_table_cells(translation_dict, table, step, newest_print_progress_threshold, print_progress_increment, count_of_relevant_paragraphs): #add doc if debugging is needed
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                count_of_relevant_paragraphs += process_paragraph_and_runs_within_it(translation_dict, paragraph, step) #add doc if debugging is needed
                newest_print_progress_threshold = indicate_progress(translation_dict, step, newest_print_progress_threshold, print_progress_increment, count_of_relevant_paragraphs)

            # Recursively process any nested tables inside the current cell
            for nested_table in cell.tables:
                process_table_cells(translation_dict, nested_table, step, newest_print_progress_threshold, print_progress_increment, count_of_relevant_paragraphs) #add doc if debugging is needed

    return count_of_relevant_paragraphs


#__________________________________________________________________________
###########################################################################
# Function to extract or swap a paragraph's runs after first consolidating the paragraph's runs
# Returns: 0 if extraction successful, 1 if swapping successful, -1 if neither extraction nor swapping operations occurred
def process_paragraph_and_runs_within_it(translation_dict, paragraph, step): #add doc if debugging is needed
    if is_relevant_paragraph(paragraph):

        # Iterate over runs in the paragraph to consolidate them
        cons_paragraph = consolidate_runs(paragraph) #add doc if debugging is needed

        if step == constants.EXTRACT:
            # Iterate over the consolidated runs in the paragraph to extract text on a consolidated-run basis
            paragraph_tagged_source_text_with_preserves = extract_runs(cons_paragraph)

            if (paragraph_tagged_source_text_with_preserves != "" 
                and not paragraph_tagged_source_text_with_preserves.isspace()
                and paragraph_tagged_source_text_with_preserves not in translation_dict
                ):
                # Add it to the translation dictionary
                translation_dict[paragraph_tagged_source_text_with_preserves] = {
                    "paragraph_tagged_translated_text_with_preserves": None
                }
            
                # Extraction: 0 if successful
                return 0

        if step == constants.SWAP:
            # Iterate over runs in the paragraph to swap text on a consolidated-run basis
            (paragraph, current_swap_count) = paragraph_level_swapper(translation_dict, cons_paragraph) #add doc if debugging is needed
            
            # Swapping: 1 if successful, 0 if failure
            return current_swap_count
    
    # Irrelevant: -1 if a repeat or empty paragraph is found, i.e. no operation occurred
    return -1

